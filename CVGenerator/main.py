from docx import Document

doc = Document()





user_name=input("What is your name?: ").strip()
user_surname=input("Enter your surname: ").strip()
phone=input("Enter your phone number: ").strip()
email=input("Enter your phone email: ").strip()
title=input("What is your professional title?: ").strip()
born_year=input("What year were you born (January 1, 2000): ").strip()
user_address=input("Enter your address: ").strip()
link=input("Your Linkedin or Github link: ").strip()
b_degree=input("Do you have a Bachelor degree?: ").strip().lower()
b_university=None
b_field=None
b_graduation=None
m_university=None
m_degree=None
m_field=None
m_graduation=None
if b_degree=='yes':
    b_university = input("What is the name of your University?: ").strip()
    b_field=input("In what field do you have a degree?: ").strip()
    b_graduation = input("What is your (expected) graduation year(Bachelors)?: ").strip()
    m_degree = input("Do you have a Masters degree?: ").strip().lower()

    if m_degree == 'yes':
        m_university = input("In which university did you complete Masters?: ").strip()
        m_field = input("In what field do you have a degree?: ").strip()
        m_graduation = input("What is your (expected) graduation year(Masters)?: ").strip()
specialization=input("What is your primary area of focus?: ").strip()
hard_skills=input("Enter your technical skills: ").strip().split()
soft_skills=input("Enter your soft skills: ").strip().split()
language=input("What languages are you fluent?: ").strip()
photo=input("Would you like to include a profile photo?: ").strip()
if photo=='yes':
    p=input("Send your photo: ")
    with open('user_photo.jpg', 'wb') as f:
        f.write(p)
experience=input('Do you have any work experience or internships to add?: ').strip().lower()
user_experience=None
if experience=='yes':
    role=input("What was your job title?: ").strip().title()
    company=input("What was the name of the Company or Organization?: ").strip().title()
    timeline=input("When did you start and end?: ").strip().title()
    user_experience={"role":None if role=="" else role, "company": None if company=="" else company.split(), "timeline": timeline}
full_name=f"{user_name.title()}{user_surname.title()}"
doc.add_heading(full_name, 0)


template_1=f"""
{title.title()}
Born: {born_year.title()}
{phone} | {email} | {link}
{user_address.title()}


{"EDUCATION" if b_degree=='yes' else ''}


{f"{b_university.title()}" if b_university else ''}
{f"Bachelors of {b_field.title()}" if b_degree=='yes' else ''}
{f"Graduation: {b_graduation}" if b_degree=='yes' else ''}

{f"{m_university.title()}" if m_university else ''}
{f"Masters of {m_field.title()}" if m_degree else ''}
{f"Graduation: {m_graduation}" if m_graduation else ''}


{"EXPERIENCE" if experience=='yes' else ''}
{f"{user_experience['role'].title()} at {user_experience['company']} — {user_experience['timeline']}"}


SKILLS

Hard skills: {', '.join(hard_skills)}
Soft skills: {', '.join(soft_skills)}


{"ADDITIONAL INFORMATION" if language else ''}
{"Languages" if language!='' else ''}: {', '.join(language)}
""".split('\n')
for i in range(len(template_1)):
    doc.add_paragraph(template_1[i])

doc.save(f'{full_name}CV.docx')



