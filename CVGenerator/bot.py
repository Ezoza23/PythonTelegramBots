from telegram import Update
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler,
    filters, ContextTypes, ConversationHandler
)
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import os

# ---------- STATES ----------
(
    NAME, SURNAME, PHONE, EMAIL, TITLE, BORN_YEAR,
    ADDRESS, LINK,
    B_DEGREE, B_UNIVERSITY, B_FIELD, B_GRAD,
    M_DEGREE, M_UNIVERSITY, M_FIELD, M_GRAD,
    SPECIALIZATION, HARD_SKILLS, SOFT_SKILLS, LANGUAGE,
    PHOTO,
    EXPERIENCE, ROLE, COMPANY, TIMELINE
) = range(25)

# ---------- HELPERS ----------
def unique_name():
    return uuid.uuid4().hex[:8]

def add_sections(doc, data):
    # ADD PHOTO
    if data.get('photo'):
        try:
            doc.add_picture(data['photo'], width=Inches(1.2))
        except Exception as e:
            print(f"Photo could not be added: {e}")

    # CONTACT
    doc.add_paragraph(f"{data.get('phone','')} | {data.get('email','')} | {data.get('link','')}")
    doc.add_paragraph(data.get('address',''))

    # PROFILE/SPECIALIZATION
    if data.get('specialization'):
        doc.add_heading("PROFILE", level=1)
        doc.add_paragraph(data.get('specialization',''))

    # EDUCATION
    if data.get('b_degree') == 'yes':
        doc.add_heading("EDUCATION", level=1)
        doc.add_paragraph(f"{data.get('b_university','').title()} — {data.get('b_field','').title()}")
        doc.add_paragraph(f"Graduation: {data.get('b_grad','')}")
        if data.get('m_degree') == 'yes':
            doc.add_paragraph(f"{data.get('m_university','').title()} — {data.get('m_field','').title()}")
            doc.add_paragraph(f"Graduation: {data.get('m_grad','')}")

    # EXPERIENCE
    if data.get('experience') == 'yes':
        doc.add_heading("EXPERIENCE", level=1)
        doc.add_paragraph(f"{data.get('role','').title()} at {data.get('company','').title()} — {data.get('timeline','')}")

    # SKILLS
    doc.add_heading("SKILLS", level=1)
    doc.add_paragraph("Hard skills: " + ", ".join(data.get('hard_skills', [])))
    doc.add_paragraph("Soft skills: " + ", ".join(data.get('soft_skills', [])))

    # LANGUAGES
    if data.get('language'):
        doc.add_heading("ADDITIONAL INFORMATION", level=1)
        doc.add_paragraph(f"Languages: {data.get('language')}")

# ---------- TEMPLATE FUNCTIONS ----------
def template1(data):
    # Classic single-column
    doc = Document()
    doc.add_heading(f"{data.get('name').upper()} {data.get('surname').upper()}", 0)
    doc.add_paragraph(data.get('title','').upper())
    add_sections(doc, data)
    file = f"{unique_name()}_t1.docx"
    doc.save(file)
    return file

def template2(data):
    # Centered elegant
    doc = Document()
    p = doc.add_paragraph(f"{data.get('name').upper()} {data.get('surname').upper()}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(data.get('title','').upper())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_sections(doc, data)
    file = f"{unique_name()}_t2.docx"
    doc.save(file)
    return file

def template3(data):
    # Two-column modern
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells

    # LEFT COLUMN
    if data.get('photo'):
        try:
            left_paragraph = left.add_paragraph()
            left_paragraph.add_run().add_picture(data['photo'], width=Inches(1.0))
        except:
            pass
    left.add_paragraph(f"{data.get('phone','')}\n{data.get('email','')}\n{data.get('link','')}")
    left.add_paragraph("Hard skills: " + ", ".join(data.get('hard_skills', [])))
    left.add_paragraph("Soft skills: " + ", ".join(data.get('soft_skills', [])))
    left.add_paragraph(f"Languages: {data.get('language','')}")

    # RIGHT COLUMN
    right.add_paragraph(f"{data.get('name').upper()} {data.get('surname').upper()}", style='Heading1')
    right.add_paragraph(data.get('title','').upper())
    if data.get('specialization'):
        right.add_paragraph("PROFILE:\n" + data.get('specialization',''))
    if data.get('b_degree') == 'yes':
        right.add_paragraph(f"Education:\n{data.get('b_university','').title()} — {data.get('b_field','').title()}")
        right.add_paragraph(f"Graduation: {data.get('b_grad','')}")
        if data.get('m_degree') == 'yes':
            right.add_paragraph(f"{data.get('m_university','').title()} — {data.get('m_field','').title()}")
            right.add_paragraph(f"Graduation: {data.get('m_grad','')}")
    if data.get('experience') == 'yes':
        right.add_paragraph(f"Experience:\n{data.get('role','').title()} at {data.get('company','').title()} — {data.get('timeline','')}")

    file = f"{unique_name()}_t3.docx"
    doc.save(file)
    return file

def template4(data):
    # Horizontal sections, modern
    doc = Document()
    doc.add_heading(f"{data.get('name').upper()} {data.get('surname').upper()}", 0)
    doc.add_heading("PROFILE", 1)
    doc.add_paragraph(data.get('specialization',''))
    doc.add_heading("EDUCATION", 1)
    if data.get('b_degree')=='yes':
        doc.add_paragraph(f"{data.get('b_university','').title()} — {data.get('b_field','').title()}")
        doc.add_paragraph(f"Graduation: {data.get('b_grad','')}")
    if data.get('m_degree')=='yes':
        doc.add_paragraph(f"{data.get('m_university','').title()} — {data.get('m_field','').title()}")
        doc.add_paragraph(f"Graduation: {data.get('m_grad','')}")
    doc.add_heading("EXPERIENCE", 1)
    if data.get('experience')=='yes':
        doc.add_paragraph(f"{data.get('role','').title()} at {data.get('company','').title()} — {data.get('timeline','')}")
    doc.add_heading("SKILLS",1)
    doc.add_paragraph("Hard skills: " + ", ".join(data.get('hard_skills',[])))
    doc.add_paragraph("Soft skills: " + ", ".join(data.get('soft_skills',[])))
    doc.add_heading("ADDITIONAL INFO",1)
    doc.add_paragraph(f"Languages: {data.get('language','')}")
    if data.get('photo'):
        try:
            doc.add_picture(data['photo'], width=Inches(1.0))
        except:
            pass
    file = f"{unique_name()}_t4.docx"
    doc.save(file)
    return file

def template5(data):
    # Minimalist creative
    doc = Document()
    doc.add_paragraph("="*40)
    doc.add_paragraph(f"{data.get('name').upper()} {data.get('surname').upper()}")
    doc.add_paragraph("="*40)
    doc.add_paragraph(f"TITLE: {data.get('title','').upper()}")
    doc.add_paragraph(f"EMAIL: {data.get('email','')}")
    doc.add_paragraph(f"PHONE: {data.get('phone','')}")
    add_sections(doc, data)
    file = f"{unique_name()}_t5.docx"
    doc.save(file)
    return file

# ---------- GENERATE AND SEND ----------
async def finish(update, context):
    data = context.user_data
    files = [
        template1(data),
        template2(data),
        template3(data),
        template4(data),
        template5(data)
    ]

    for i, file in enumerate(files,1):
        with open(file,'rb') as f:
            await update.message.reply_document(f, filename=f"Template_{i}.docx")
        os.remove(file)
    await update.message.reply_text("✅ I sent 5 different CV templates. Choose the one you like!")
    return ConversationHandler.END

# ---------- CONVERSATION HANDLERS ----------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("What is your name?")
    return NAME

async def name(update, context): context.user_data['name']=update.message.text.strip(); await update.message.reply_text("Enter your surname:"); return SURNAME
async def surname(update, context): context.user_data['surname']=update.message.text.strip(); await update.message.reply_text("Enter your phone number:"); return PHONE
async def phone(update, context): context.user_data['phone']=update.message.text.strip(); await update.message.reply_text("Enter your email:"); return EMAIL
async def email(update, context): context.user_data['email']=update.message.text.strip(); await update.message.reply_text("What is your professional title?"); return TITLE
async def title(update, context): context.user_data['title']=update.message.text.strip(); await update.message.reply_text("What year were you born?"); return BORN_YEAR
async def born(update, context): context.user_data['born']=update.message.text.strip(); await update.message.reply_text("Enter your address:"); return ADDRESS
async def address(update, context): context.user_data['address']=update.message.text.strip(); await update.message.reply_text("Your Linkedin or Github link:"); return LINK
async def link(update, context): context.user_data['link']=update.message.text.strip(); await update.message.reply_text("Do you have a Bachelor degree? (yes/no)"); return B_DEGREE
async def b_degree(update, context):
    ans=update.message.text.strip().lower(); context.user_data['b_degree']=ans
    if ans=='yes': await update.message.reply_text("What is the name of your University?"); return B_UNIVERSITY
    else: await update.message.reply_text("What is your primary area of focus?"); return SPECIALIZATION
async def b_university(update, context): context.user_data['b_university']=update.message.text.strip(); await update.message.reply_text("In what field do you have a degree?"); return B_FIELD
async def b_field(update, context): context.user_data['b_field']=update.message.text.strip(); await update.message.reply_text("What is your graduation year (Bachelors)?"); return B_GRAD
async def b_grad(update, context): context.user_data['b_grad']=update.message.text.strip(); await update.message.reply_text("Do you have a Masters degree? (yes/no)"); return M_DEGREE
async def m_degree(update, context):
    ans=update.message.text.strip().lower(); context.user_data['m_degree']=ans
    if ans=='yes': await update.message.reply_text("In which university did you complete Masters?"); return M_UNIVERSITY
    else: await update.message.reply_text("What is your primary area of focus?"); return SPECIALIZATION
async def m_university(update, context): context.user_data['m_university']=update.message.text.strip(); await update.message.reply_text("In what field do you have a degree?"); return M_FIELD
async def m_field(update, context): context.user_data['m_field']=update.message.text.strip(); await update.message.reply_text("What is your graduation year (Masters)?"); return M_GRAD
async def m_grad(update, context): context.user_data['m_grad']=update.message.text.strip(); await update.message.reply_text("What is your primary area of focus?"); return SPECIALIZATION
async def specialization(update, context): context.user_data['specialization']=update.message.text.strip(); await update.message.reply_text("Enter your technical skills (space separated):"); return HARD_SKILLS
async def hard_skills(update, context): context.user_data['hard_skills']=update.message.text.split(); await update.message.reply_text("Enter your soft skills (space separated):"); return SOFT_SKILLS
async def soft_skills(update, context): context.user_data['soft_skills']=update.message.text.split(); await update.message.reply_text("What languages are you fluent in?"); return LANGUAGE
async def language(update, context): context.user_data['language']=update.message.text.strip(); await update.message.reply_text("Send your profile photo or type 'skip'"); return PHOTO
async def photo(update, context):
    if update.message.photo:
        photo_file = await update.message.photo[-1].get_file()
        path = f"{uuid.uuid4().hex}.jpg"
        await photo_file.download_to_drive(path)
        context.user_data['photo'] = path
    else:
        context.user_data['photo'] = None
    await update.message.reply_text("Do you have any work experience? (yes/no)"); return EXPERIENCE
async def experience(update, context):
    ans=update.message.text.strip().lower(); context.user_data['experience']=ans
    if ans=='yes': await update.message.reply_text("What was your job title?"); return ROLE
    else: return await finish(update, context)
async def role(update, context): context.user_data['role']=update.message.text.strip(); await update.message.reply_text("Company name?"); return COMPANY
async def company(update, context): context.user_data['company']=update.message.text.strip(); await update.message.reply_text("Timeline (start - end):"); return TIMELINE
async def timeline(update, context): context.user_data['timeline']=update.message.text.strip(); return await finish(update, context)

# ---------- MAIN ----------
app = ApplicationBuilder().token("8639212152:AAGhwFkaFFg6OAVdyW1hdkQbvAE5R85hJjg").build()

conv_handler = ConversationHandler(
    entry_points=[CommandHandler("start", start)],
    states={
        NAME:[MessageHandler(filters.TEXT, name)],
        SURNAME:[MessageHandler(filters.TEXT, surname)],
        PHONE:[MessageHandler(filters.TEXT, phone)],
        EMAIL:[MessageHandler(filters.TEXT, email)],
        TITLE:[MessageHandler(filters.TEXT, title)],
        BORN_YEAR:[MessageHandler(filters.TEXT, born)],
        ADDRESS:[MessageHandler(filters.TEXT, address)],
        LINK:[MessageHandler(filters.TEXT, link)],
        B_DEGREE:[MessageHandler(filters.TEXT, b_degree)],
        B_UNIVERSITY:[MessageHandler(filters.TEXT, b_university)],
        B_FIELD:[MessageHandler(filters.TEXT, b_field)],
        B_GRAD:[MessageHandler(filters.TEXT, b_grad)],
        M_DEGREE:[MessageHandler(filters.TEXT, m_degree)],
        M_UNIVERSITY:[MessageHandler(filters.TEXT, m_university)],
        M_FIELD:[MessageHandler(filters.TEXT, m_field)],
        M_GRAD:[MessageHandler(filters.TEXT, m_grad)],
        SPECIALIZATION:[MessageHandler(filters.TEXT, specialization)],
        HARD_SKILLS:[MessageHandler(filters.TEXT, hard_skills)],
        SOFT_SKILLS:[MessageHandler(filters.TEXT, soft_skills)],
        LANGUAGE:[MessageHandler(filters.TEXT, language)],
        PHOTO:[MessageHandler(filters.PHOTO | filters.TEXT, photo)],
        EXPERIENCE:[MessageHandler(filters.TEXT, experience)],
        ROLE:[MessageHandler(filters.TEXT, role)],
        COMPANY:[MessageHandler(filters.TEXT, company)],
        TIMELINE:[MessageHandler(filters.TEXT, timeline)],
    },
    fallbacks=[]
)

app.add_handler(conv_handler)
app.run_polling()